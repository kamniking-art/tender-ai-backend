from __future__ import annotations

from datetime import UTC, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models import Company
from app.tender_analysis.model import TenderAnalysis
from app.tender_analysis.service import get_analysis_scoped
from app.tender_decisions.service import get_decision_scoped
from app.tender_documents.model import TenderDocument
from app.tender_documents.service import create_document_from_bytes
from app.tenders.model import Tender
from app.tenders.service import get_tender_by_id_scoped

PACKAGE_DIR_NAME = "package_v1"
DEFAULT_CHECKLIST = [
    "Заявка (форма участника)",
    "Учредительные документы / устав",
    "Выписка ЕГРЮЛ (актуальная)",
    "Документы полномочий подписанта (приказ/доверенность)",
    "Декларации соответствия требованиям",
]

REQUIRED_PROFILE_FIELDS = [
    "legal_name",
    "inn",
    "legal_address",
    "director_name",
    "phone",
    "email",
]


class DocumentModuleError(Exception):
    pass


class DocumentModuleNotFoundError(DocumentModuleError):
    pass


class DocumentModuleConflictError(DocumentModuleError):
    pass


class DocumentModuleValidationError(DocumentModuleError):
    def __init__(self, detail: str, missing_fields: list[str] | None = None) -> None:
        super().__init__(detail)
        self.missing_fields = missing_fields or []


class GeneratedFileResult:
    def __init__(self, document_id: UUID, filename: str) -> None:
        self.document_id = document_id
        self.filename = filename


class PackageState:
    def __init__(self, files: list[TenderDocument], checklist: list[str]) -> None:
        self.files = files
        self.checklist = checklist


def _package_prefix(company_id: UUID, tender_id: UUID) -> str:
    return (Path(settings.documents_subdir) / str(company_id) / str(tender_id) / PACKAGE_DIR_NAME).as_posix()


def _format_decimal(value: Decimal | None) -> str:
    if value is None:
        return ""
    normalized = value.quantize(Decimal("0.01"))
    return f"{normalized}"


def _safe_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value)


def _extract_profile(company: Company) -> dict:
    return company.profile if isinstance(company.profile, dict) else {}


def _missing_profile_fields(profile: dict) -> list[str]:
    return [field for field in REQUIRED_PROFILE_FIELDS if not profile.get(field)]


def _detect_security_flags(analysis: TenderAnalysis | None, decision) -> tuple[bool, bool]:
    extracted = (analysis.requirements or {}).get("extracted_v1") if analysis else None
    extracted = extracted if isinstance(extracted, dict) else {}

    has_bid_security = bool(
        extracted.get("bid_security_required")
        or extracted.get("bid_security_pct") is not None
        or extracted.get("bid_security_amount") is not None
        or (decision and decision.need_bid_security)
    )
    has_contract_security = bool(
        extracted.get("contract_security_required")
        or extracted.get("contract_security_pct") is not None
        or extracted.get("contract_security_amount") is not None
        or (decision and decision.need_contract_security)
    )

    flags = analysis.risk_flags if analysis and isinstance(analysis.risk_flags, list) else []
    flag_codes = {item.get("code") for item in flags if isinstance(item, dict)}
    if "high_bid_security" in flag_codes:
        has_bid_security = True
    if "high_contract_security" in flag_codes:
        has_contract_security = True

    return has_bid_security, has_contract_security


def build_checklist(analysis: TenderAnalysis | None, decision) -> list[str]:
    checklist = list(DEFAULT_CHECKLIST)

    extracted = (analysis.requirements or {}).get("extracted_v1") if analysis else None
    extracted = extracted if isinstance(extracted, dict) else {}
    qualification_requirements = extracted.get("qualification_requirements")
    if isinstance(qualification_requirements, list) and qualification_requirements:
        checklist.append("Подтверждающие документы по квалификации (опыт/договоры/акты)")

    has_bid_security, has_contract_security = _detect_security_flags(analysis, decision)
    if has_bid_security:
        checklist.append("Документ обеспечения заявки (БГ/платежное поручение)")
    if has_contract_security:
        checklist.append("Документ обеспечения исполнения контракта (БГ/платежное поручение)")

    return checklist


def render_docx_template(template_path: Path, context_dict: dict[str, str]) -> bytes:
    document = Document(str(template_path))

    def replace_text(text: str) -> str:
        rendered = text
        for key, value in context_dict.items():
            rendered = rendered.replace(f"{{{{{key}}}}}", value)
        return rendered

    for paragraph in document.paragraphs:
        replaced = replace_text(paragraph.text)
        if replaced != paragraph.text:
            paragraph.text = replaced

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced = replace_text(paragraph.text)
                    if replaced != paragraph.text:
                        paragraph.text = replaced

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


async def _get_company_scoped(db: AsyncSession, company_id: UUID) -> Company | None:
    return await db.scalar(select(Company).where(Company.id == company_id))


async def _list_package_documents(db: AsyncSession, company_id: UUID, tender_id: UUID) -> list[TenderDocument]:
    prefix = f"{_package_prefix(company_id, tender_id)}/"
    stmt = (
        select(TenderDocument)
        .where(
            TenderDocument.company_id == company_id,
            TenderDocument.tender_id == tender_id,
            TenderDocument.storage_path.like(f"{prefix}%"),
        )
        .order_by(TenderDocument.uploaded_at.desc())
    )
    return list((await db.scalars(stmt)).all())


async def _delete_package_documents(db: AsyncSession, documents: list[TenderDocument]) -> None:
    for doc in documents:
        file_path = Path(settings.storage_root) / doc.storage_path
        try:
            if file_path.exists() and file_path.is_file():
                file_path.unlink()
        except OSError:
            pass
        await db.delete(doc)


def _template_context(company_profile: dict, tender: Tender, analysis: TenderAnalysis | None, decision) -> dict[str, str]:
    extracted = (analysis.requirements or {}).get("extracted_v1") if analysis else None
    extracted = extracted if isinstance(extracted, dict) else {}

    deadline = extracted.get("submission_deadline_at") or tender.submission_deadline
    bid_security = extracted.get("bid_security_amount") or extracted.get("bid_security_pct")
    contract_security = extracted.get("contract_security_amount") or extracted.get("contract_security_pct")

    return {
        "LEGAL_NAME": _safe_text(company_profile.get("legal_name")),
        "INN": _safe_text(company_profile.get("inn")),
        "ADDRESS": _safe_text(company_profile.get("legal_address")),
        "DIRECTOR": _safe_text(company_profile.get("director_name")),
        "PHONE": _safe_text(company_profile.get("phone")),
        "EMAIL": _safe_text(company_profile.get("email")),
        "TENDER_SUBJECT": _safe_text(tender.title),
        "TENDER_NMCK": _format_decimal(tender.nmck),
        "DEADLINE_AT": _safe_text(deadline),
        "BID_SECURITY": _safe_text(bid_security),
        "CONTRACT_SECURITY": _safe_text(contract_security),
    }


def _build_package_relative_paths(company_id: UUID, tender_id: UUID) -> dict[str, str]:
    base = Path(settings.documents_subdir) / str(company_id) / str(tender_id) / PACKAGE_DIR_NAME
    return {
        "checklist": (base / "checklist.txt").as_posix(),
        "application": (base / "drafts" / "application_draft.docx").as_posix(),
        "cover": (base / "drafts" / "cover_letter.docx").as_posix(),
    }


async def generate_package_for_tender(
    db: AsyncSession,
    *,
    company_id: UUID,
    tender_id: UUID,
    user_id: UUID,
    force: bool,
) -> tuple[list[GeneratedFileResult], list[str]]:
    tender = await get_tender_by_id_scoped(db, company_id, tender_id)
    if tender is None:
        raise DocumentModuleNotFoundError("Tender not found")

    decision = await get_decision_scoped(db, company_id, tender_id)
    if decision is None or decision.recommendation != "go":
        raise DocumentModuleConflictError("Decision is not GO")

    company = await _get_company_scoped(db, company_id)
    if company is None:
        raise DocumentModuleNotFoundError("Company not found")

    profile = _extract_profile(company)
    missing_fields = _missing_profile_fields(profile)
    if missing_fields:
        raise DocumentModuleValidationError("Company profile is incomplete", missing_fields=missing_fields)

    existing_docs = await _list_package_documents(db, company_id, tender_id)
    if existing_docs and not force:
        raise DocumentModuleConflictError("Package already generated")

    if existing_docs and force:
        await _delete_package_documents(db, existing_docs)
        await db.commit()

    analysis = await get_analysis_scoped(db, company_id, tender_id)
    checklist = build_checklist(analysis, decision)
    relative_paths = _build_package_relative_paths(company_id, tender_id)

    checklist_content = "\n".join(f"- {item}" for item in checklist) + "\n"

    template_dir = Path(__file__).resolve().parent / "templates"
    context = _template_context(profile, tender, analysis, decision)
    application_doc = render_docx_template(template_dir / "application_draft_v1.docx", context)
    cover_doc = render_docx_template(template_dir / "cover_letter_v1.docx", context)

    generated_docs: list[TenderDocument] = []
    generated_docs.append(
        await create_document_from_bytes(
            db,
            company_id=company_id,
            tender_id=tender_id,
            uploaded_by=user_id,
            file_name="checklist.txt",
            content=checklist_content.encode("utf-8"),
            content_type="text/plain",
            doc_type="package",
            relative_path_override=relative_paths["checklist"],
        )
    )
    generated_docs.append(
        await create_document_from_bytes(
            db,
            company_id=company_id,
            tender_id=tender_id,
            uploaded_by=user_id,
            file_name="application_draft.docx",
            content=application_doc,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            doc_type="draft",
            relative_path_override=relative_paths["application"],
        )
    )
    generated_docs.append(
        await create_document_from_bytes(
            db,
            company_id=company_id,
            tender_id=tender_id,
            uploaded_by=user_id,
            file_name="cover_letter.docx",
            content=cover_doc,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            doc_type="draft",
            relative_path_override=relative_paths["cover"],
        )
    )

    generated_at = datetime.now(UTC).isoformat().replace("+00:00", "Z")
    if analysis is None:
        analysis = TenderAnalysis(
            company_id=company_id,
            tender_id=tender_id,
            status="draft",
            requirements={"documents_v1": {"checklist": checklist, "generated_at": generated_at}},
            missing_docs=[],
            risk_flags=[],
            summary=None,
            created_by=user_id,
            updated_by=user_id,
        )
        db.add(analysis)
    else:
        requirements = dict(analysis.requirements or {})
        requirements["documents_v1"] = {"checklist": checklist, "generated_at": generated_at}
        analysis.requirements = requirements
        analysis.updated_by = user_id

    await db.commit()

    return [GeneratedFileResult(document_id=item.id, filename=item.file_name) for item in generated_docs], checklist


async def get_package_for_tender(
    db: AsyncSession,
    *,
    company_id: UUID,
    tender_id: UUID,
) -> PackageState:
    tender = await get_tender_by_id_scoped(db, company_id, tender_id)
    if tender is None:
        raise DocumentModuleNotFoundError("Tender not found")

    files = await _list_package_documents(db, company_id, tender_id)
    analysis = await get_analysis_scoped(db, company_id, tender_id)

    checklist: list[str] = []
    if analysis and isinstance(analysis.requirements, dict):
        docs_payload = analysis.requirements.get("documents_v1")
        if isinstance(docs_payload, dict) and isinstance(docs_payload.get("checklist"), list):
            checklist = [str(item) for item in docs_payload.get("checklist", [])]

    return PackageState(files=files, checklist=checklist)
