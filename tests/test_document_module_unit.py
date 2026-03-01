from __future__ import annotations

import tempfile
from pathlib import Path
from unittest import TestCase

from docx import Document

from app.document_module.service import render_docx_template


class DocumentModuleUnitTests(TestCase):
    def test_render_docx_template_replaces_paragraphs_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"

            doc = Document()
            doc.add_paragraph("Hello {{LEGAL_NAME}}")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "INN={{INN}}"
            doc.save(template_path)

            rendered = render_docx_template(
                template_path,
                {
                    "LEGAL_NAME": "ООО Тест",
                    "INN": "1234567890",
                },
            )
            output_path.write_bytes(rendered)

            rendered_doc = Document(str(output_path))
            self.assertEqual(rendered_doc.paragraphs[0].text, "Hello ООО Тест")
            self.assertEqual(rendered_doc.tables[0].cell(0, 0).text, "INN=1234567890")
